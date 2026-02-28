from __future__ import annotations

import re
import tempfile
from datetime import datetime, timezone
from typing import Iterable

from bson import ObjectId
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse

from app.core.db import get_db
from app.models.tailor import (
    JobIngestIn,
    JobIngestOut,
    ExtractedSkill,
    TailorPreviewIn,
    TailoredResumeOut,
    ResumeSection,
)
from app.utils.mongo import oid_str

router = APIRouter()

def now_utc():
    return datetime.now(timezone.utc)

def _tokenize_keywords(text: str) -> list[str]:
    # lightweight keywords: alphanum words len>=4 excluding very common
    stop = {"with","from","that","this","have","will","your","able","work","team","role","must","plus","also","using","used","into","over","such","they","their","them","than","then","only","when","where","what","were","been","being","more","less","some","many","each","make","made"}
    words = re.findall(r"[A-Za-z][A-Za-z0-9+.#-]{3,}", text.lower())
    out = []
    seen = set()
    for w in words:
        if w in stop:
            continue
        if w in seen:
            continue
        seen.add(w)
        out.append(w)
        if len(out) >= 25:
            break
    return out

async def _load_skill_catalog(db) -> list[dict]:
    cursor = db["skills"].find({}, {"name": 1, "aliases": 1})
    return await cursor.to_list(length=5000)

def _match_skills(job_text: str, skills: Iterable[dict]) -> list[ExtractedSkill]:
    text = job_text.lower()
    matches: dict[str, ExtractedSkill] = {}

    def bump(sid: str, name: str, matched_on: str):
        key = sid
        if key not in matches:
            matches[key] = ExtractedSkill(skill_id=sid, skill_name=name, matched_on=matched_on, count=1)
        else:
            matches[key].count += 1

    for s in skills:
        sid = str(s["_id"])
        name = (s.get("name") or "").strip()
        if not name:
            continue

        # word-boundary match for names; allow special tokens like C++ / C# via loose matching
        n = name.lower()
        if len(n) >= 2:
            if re.search(rf"(?<![A-Za-z0-9]){re.escape(n)}(?![A-Za-z0-9])", text):
                bump(sid, name, "name")

        for a in s.get("aliases", []) or []:
            a = (a or "").strip()
            if not a:
                continue
            al = a.lower()
            if re.search(rf"(?<![A-Za-z0-9]){re.escape(al)}(?![A-Za-z0-9])", text):
                bump(sid, name, "alias")

    # sort by count desc then name
    return sorted(matches.values(), key=lambda x: (-x.count, x.skill_name.lower()))

async def _load_user_items(db, user_id: str) -> list[dict]:
    # prefer unified portfolio_items
    items = await db["portfolio_items"].find({"user_id": user_id}).to_list(length=2000)
    # fallback: projects as items (if you haven't migrated yet)
    if not items:
        projs = await db["projects"].find({"user_id": user_id}).to_list(length=500)
        for p in projs:
            items.append({
                "_id": p["_id"],
                "type": "project",
                "title": p.get("title", ""),
                "org": None,
                "summary": p.get("description"),
                "bullets": [],
                "links": [],
                "skill_ids": [],
                "priority": 0,
                "updated_at": p.get("updated_at"),
                "created_at": p.get("created_at"),
            })
    return items

def _score_item(item: dict, job_skill_ids: set[str], keywords: set[str]) -> float:
    sids = set(item.get("skill_ids", []) or [])
    overlap = len(sids & job_skill_ids)
    pri = float(item.get("priority", 0) or 0)

    # keyword overlap from title + summary + bullets
    txt = " ".join(
        [item.get("title",""), item.get("summary","")] + (item.get("bullets", []) or [])
    ).lower()
    kw_hit = sum(1 for k in keywords if k in txt)

    return overlap * 5.0 + kw_hit * 1.0 + pri * 0.25

def _render_plain_text(sections: list[ResumeSection]) -> str:
    lines: list[str] = []
    for sec in sections:
        lines.append(sec.title.upper())
        for ln in sec.lines:
            lines.append(ln)
        lines.append("")
    return "\n".join(lines).strip() + "\n"

@router.post("/job/ingest", response_model=JobIngestOut)
async def ingest_job(payload: JobIngestIn):
    db = get_db()
    skills = await _load_skill_catalog(db)
    extracted = _match_skills(payload.text, skills)
    keywords = _tokenize_keywords(payload.text)

    now = now_utc()
    doc = payload.model_dump()
    doc["extracted_skills"] = [e.model_dump() for e in extracted[:50]]
    doc["keywords"] = keywords
    doc["created_at"] = now
    # store full text; keep preview for response
    res = await db["job_ingests"].insert_one(doc)

    preview = payload.text.strip().replace("\n", " ")
    if len(preview) > 220:
        preview = preview[:220] + "..."

    return JobIngestOut(
        id=oid_str(res.inserted_id),
        user_id=payload.user_id,
        title=payload.title,
        company=payload.company,
        location=payload.location,
        text_preview=preview,
        extracted_skills=extracted[:25],
        keywords=keywords,
        created_at=now,
    )

@router.post("/preview", response_model=TailoredResumeOut)
async def preview_tailored_resume(payload: TailorPreviewIn):
    db = get_db()

    job_text = payload.job_text
    job_id = payload.job_id

    extracted: list[ExtractedSkill] = []
    keywords: list[str] = []

    if job_id:
        try:
            job_oid = ObjectId(job_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid job_id")

        job_doc = await db["job_ingests"].find_one({"_id": job_oid, "user_id": payload.user_id})
        if not job_doc:
            raise HTTPException(status_code=404, detail="Job ingest not found for user_id")
        job_text = job_doc.get("text", "")
        extracted = [ExtractedSkill(**e) for e in (job_doc.get("extracted_skills") or [])]
        keywords = job_doc.get("keywords") or []
    else:
        if not job_text or len(job_text) < 50:
            raise HTTPException(status_code=400, detail="Provide job_id or job_text (>=50 chars)")
        skills = await _load_skill_catalog(db)
        extracted = _match_skills(job_text, skills)
        keywords = _tokenize_keywords(job_text)

    job_skill_ids = {e.skill_id for e in extracted[:50]}
    keyword_set = set(keywords)

    # load user portfolio items
    items = await _load_user_items(db, payload.user_id)

    scored = []
    for it in items:
        scored.append(( _score_item(it, job_skill_ids, keyword_set), it))
    scored.sort(key=lambda x: x[0], reverse=True)

    selected_items = [it for score, it in scored[: payload.max_items] if score > 0] or [it for score, it in scored[: payload.max_items]]
    selected_item_ids = [oid_str(it["_id"]) for it in selected_items if "_id" in it]

    # Select skills: prioritize skills that appear in both job and user's confirmed skills if available
    confirmed_skill_ids: set[str] = set()
    conf = await db["resume_skill_confirmations"].find_one({"user_id": payload.user_id}, sort=[("created_at", -1)])
    if conf:
        for c in conf.get("confirmed", []) or []:
            sid = c.get("skill_id")
            if sid is None:
                continue
            confirmed_skill_ids.add(str(sid))

    selected_skill_ids = [sid for sid in job_skill_ids if (sid in confirmed_skill_ids)]
    if len(selected_skill_ids) < 10:
        # backfill by job skills
        for e in extracted:
            if e.skill_id not in selected_skill_ids:
                selected_skill_ids.append(e.skill_id)
            if len(selected_skill_ids) >= 15:
                break

    # Resolve skill names for display
    skill_name_by_id: dict[str, str] = {}
    if selected_skill_ids:
        oids = []
        for sid in selected_skill_ids:
            try:
                oids.append(ObjectId(sid))
            except Exception:
                continue
        if oids:
            docs = await db["skills"].find({"_id": {"$in": oids}}, {"name": 1}).to_list(length=200)
            for d in docs:
                skill_name_by_id[oid_str(d["_id"])] = d.get("name", "")

    skill_line = ", ".join([skill_name_by_id.get(s, s) for s in selected_skill_ids if skill_name_by_id.get(s, s)])[:250]

    # Compose sections
    sections: list[ResumeSection] = []

    # Summary
    sections.append(
        ResumeSection(
            title="Summary",
            lines=[
                "Software / ML-focused builder with hands-on experience delivering projects end-to-end (API, data, and deployment).",
                f"Targeting this role by emphasizing: {skill_line}" if skill_line else "Targeting this role with a focus on the requirements and deliverables described in the posting.",
            ],
        )
    )

    # Skills
    if selected_skill_ids:
        sections.append(
            ResumeSection(
                title="Skills",
                lines=[
                    ", ".join([skill_name_by_id.get(s, s) for s in selected_skill_ids if skill_name_by_id.get(s, s)]) or skill_line
                ],
            )
        )

    # Portfolio
    proj_lines: list[str] = []
    for it in selected_items:
        title = it.get("title", "Untitled")
        org = it.get("org")
        dates = ""
        if it.get("date_start") or it.get("date_end"):
            dates = f" ({it.get('date_start','')}-{it.get('date_end','')})".replace(" -", "-").replace("-)", ")")
        header = f"{title}{' — ' + org if org else ''}{dates}".strip()
        proj_lines.append(header)

        bullets = (it.get("bullets") or [])[: payload.max_bullets_per_item]
        if bullets:
            for b in bullets:
                proj_lines.append(f"- {b}")
        else:
            summ = it.get("summary") or ""
            if summ:
                proj_lines.append(f"- {summ}")
            else:
                proj_lines.append("- Relevant portfolio item selected based on skills/keywords overlap with the job posting.")
        links = it.get("links") or []
        if links:
            proj_lines.append(f"- Links: {', '.join(links[:3])}")

        proj_lines.append("")  # spacing

    if proj_lines:
        sections.append(ResumeSection(title="Relevant Work", lines=[ln for ln in proj_lines if ln != "" ]))

    # Store tailored resume record
    now = now_utc()
    record = {
        "user_id": payload.user_id,
        "job_id": job_id,
        "template": payload.template,
        "selected_skill_ids": selected_skill_ids,
        "selected_item_ids": selected_item_ids,
        "sections": [s.model_dump() for s in sections],
        "plain_text": _render_plain_text(sections),
        "created_at": now,
    }
    res = await db["tailored_resumes"].insert_one(record)

    return TailoredResumeOut(id=oid_str(res.inserted_id), **record)

def _docx_from_sections(sections: list[ResumeSection], out_path: str):
    from docx import Document
    doc = Document()
    for sec in sections:
        doc.add_heading(sec.title, level=2)
        for ln in sec.lines:
            if ln.startswith("- "):
                doc.add_paragraph(ln[2:], style="List Bullet")
            else:
                doc.add_paragraph(ln)
    doc.save(out_path)

def _pdf_from_sections(sections: list[ResumeSection], out_path: str):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(out_path, pagesize=LETTER)
    width, height = LETTER
    x = 54
    y = height - 54
    line_h = 14

    def draw_line(text: str, bold: bool = False):
        nonlocal y
        if y < 54:
            c.showPage()
            y = height - 54
        if bold:
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 11)
        c.drawString(x, y, text[:110])
        y -= line_h

    for sec in sections:
        draw_line(sec.title.upper(), bold=True)
        for ln in sec.lines:
            draw_line(ln)
        draw_line("")

    c.save()

@router.get("/{tailored_id}/export/docx")
async def export_docx(tailored_id: str):
    db = get_db()
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    d = await db["tailored_resumes"].find_one({"_id": oid})
    if not d:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    sections = [ResumeSection(**s) for s in d.get("sections", [])]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    _docx_from_sections(sections, tmp.name)

    filename = f"tailored_resume_{tailored_id}.docx"
    return FileResponse(tmp.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

@router.get("/{tailored_id}/export/pdf")
async def export_pdf(tailored_id: str):
    db = get_db()
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    d = await db["tailored_resumes"].find_one({"_id": oid})
    if not d:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    sections = [ResumeSection(**s) for s in d.get("sections", [])]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp.close()
    _pdf_from_sections(sections, tmp.name)

    filename = f"tailored_resume_{tailored_id}.pdf"
    return FileResponse(tmp.name, media_type="application/pdf", filename=filename)
